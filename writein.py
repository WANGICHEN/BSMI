from docx import Document
from glob import glob
import zipfile
import io
import os
from copy import deepcopy
import requests
from io import BytesIO 
from docx.shared import RGBColor
import table_format



BLACK = RGBColor(0, 0, 0)

def _all_paras(doc):
    for p in doc.paragraphs: 
        yield p
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs: 
                    yield p
                for tt in c.tables:            # 巢狀表格
                    for rr in tt.rows:
                        for cc in rr.cells:
                            for p in cc.paragraphs: 
                                yield p

def _copy_rpr(dst_run, src_run):
    """把來源 run 的樣式複製到新 run（不直接賦值，避免 lxml parent 衝突）。"""
    src = src_run._element.rPr
    if src is None: 
        return
    el = dst_run._element
    if el.rPr is not None:
        el.remove(el.rPr)
    el.insert(0, deepcopy(src))

def write_doc(doc, mapping, force_black=True, prefer="first"):
    """
    mapping: 例如 {"{name}": "王小姐", "{date}": "2025/09/24"}
    force_black: True=替換文字設為黑色；False=沿用來源樣式顏色
    prefer: "first"=沿用 placeholder 起點 run 樣式；"last"=沿用終點 run 樣式
    """
    # # 先長字串優先，避免互相覆蓋
    items = sorted(mapping.items(), key=lambda kv: len(kv[0]), reverse=True)

    for para in _all_paras(doc):
        runs = list(para.runs)
        if not runs:
            continue

        # 攤平成整段文字與 run 對應區間
        full = "".join(r.text for r in runs)
        if not full:
            continue
        spans, s = [], 0
        for r in runs:
            e = s + len(r.text)
            spans.append((s, e, r))
            s = e

        # 找不重疊的所有匹配
        used = [False] * len(full)
        hits = []  # (start, end, value, base_run)
        for ph, val in items:
            i, L = 0, len(ph)
            while True:
                j = full.find(ph, i)
                if j < 0:
                    break
                k = j + L
                if all(not used[t] for t in range(j, k)):
                    involved = [r for (a, b, r) in spans if not (b <= j or a >= k)]
                    base = involved[0] if prefer == "first" else involved[-1]
                    hits.append((j, k, str(val), base))
                    for t in range(j, k):
                        used[t] = True
                i = k
        if not hits:
            continue
        hits.sort(key=lambda x: x[0])

        # 生成 new_full（只改文字，不動 run 結構）
        parts = []
        cur = 0
        for a, b, val, _base in hits:
            if cur < a:
                parts.append(full[cur:a])
            parts.append(val)
            cur = b
        if cur < len(full):
            parts.append(full[cur:])
        new_full = "".join(parts)

        if new_full == full:
            return False

        # 將 new_full 分配回原 runs：
        # 1) 先取得各 run 原本長度
        orig_lens = [len(r.text) for r in runs]

        # 2) 逐 run 填回；最後一個 run 承接剩餘（避免新增 run）
        pos = 0
        for idx, (r, L) in enumerate(zip(runs, orig_lens)):
            if idx < len(runs) - 1:
                r.text = new_full[pos:pos + L]
                pos += L
            else:
                r.text = new_full[pos:]  # 最後一個吃剩下全部

        # 3) 若你需要 force_black：只能「統一整段」或「只改特定 run」
        #    這裡採最安全策略：整段所有 runs 都變黑（避免新增 run）
        if force_black:
            for r in runs:
                r.font.color.rgb = BLACK


    return doc
    
def create_zip(file_dict):
    # file_dict: {"檔名.txt": b"檔案內容", ...}
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for fname, content in file_dict.items():
            zf.writestr(fname, content)
    buffer.seek(0)
    return buffer

def run_BSMI_doc(info):
    print("run BSMI doc")
    files = {}
    fs = [["00_08.docx", "https://z28856673-my.sharepoint.com/:w:/g/personal/itek_project_i-tek_com_tw/IQDXZ0DqrlnERZxIO01r8tq9ATwMWaJ_9AIFBDEfMMoEznw?e=8VRKis"],
         ["00_99.docx", "https://z28856673-my.sharepoint.com/:w:/g/personal/itek_project_i-tek_com_tw/IQD8OaacyRyKSK2zV3fSJRaYARSDVysOP3VNWkiGrMyo8EA?e=YcdpwT"],
         ["02_01.docx", "https://z28856673-my.sharepoint.com/:w:/g/personal/itek_project_i-tek_com_tw/IQAJNLIOp4A8R4TxNKTvPeEOAWZ5DynNGNamHoQ8L2D0AnU?e=5uweBG"],
         ["07_01.docx", "https://z28856673-my.sharepoint.com/:w:/g/personal/itek_project_i-tek_com_tw/IQAClYH5s1IwRZ53rdymbHgnAXi7rDTqgM8mCeMYs7rCpv8?e=1jofbd"],
         ["外箱標示.docx", "https://z28856673-my.sharepoint.com/:w:/g/personal/itek_project_i-tek_com_tw/IQAmL6jJKhvTQZas0kCUBy1oAfTaAlkCyYj1pvpWFVP6RG0?e=QumwfH"]]
    for f_name, f in fs:
        information = info.copy()
        if f_name in ["00_08.docx", "外箱標示.docx", "02_01.docx"]:
            information["{series}"] = ", " + information["{series}"]
        download_url = f + ("&download=1" if "?" in f else "?download=1")



        headers = {"User-Agent": "Mozilla/5.0"}
        r = requests.get(download_url, headers=headers, allow_redirects=True, timeout=30)
        r.raise_for_status()  # 403/404 會在這裡丟錯

        doc = write_doc(Document(BytesIO(r.content)), information)
        doc = table_format.set_format(f_name, doc)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        files[f_name] = buf.read()


    zip_buffer = create_zip(files)


    return zip_buffer

























